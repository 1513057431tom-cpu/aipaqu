from __future__ import annotations

from datetime import date, datetime
from io import BytesIO

from docx import Document
from fastapi import APIRouter, Depends, Response, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.core.auth import User, get_current_user
from app.core.errors import api_error
from app.core.reports import (
    DuplicateDailyReportError,
    PeriodInputIncompleteError,
    Report,
    ReportPeriod,
    ReportService,
    ReportVersion,
)
from app.core.stores import catalog_store, internal_data_store, monitoring_store, recommendation_store, report_store

router = APIRouter(prefix="/api/v1/reports", tags=["reports"])


class VersionResponse(BaseModel):
    id: str
    version: int
    markdown: str
    contentDigest: str
    changeSource: str
    createdBy: str
    createdAt: datetime


class ReportResponse(BaseModel):
    id: str
    title: str
    reportPeriod: str
    inputMode: str
    periodStart: date
    periodEnd: date
    status: str
    inputSnapshotDates: list[date]
    approvedBy: str | None
    approvedAt: datetime | None
    currentVersion: VersionResponse
    createdAt: datetime
    updatedAt: datetime


class CreateReportRequest(BaseModel):
    reportPeriod: ReportPeriod
    periodStart: date
    periodEnd: date
    title: str = Field(min_length=1, max_length=200)


class SaveVersionRequest(BaseModel):
    markdown: str = Field(min_length=1, max_length=500_000)


class ReportListResponse(BaseModel):
    data: list[ReportResponse]
    pagination: Pagination


class Pagination(BaseModel):
    page: int
    pageSize: int
    totalItems: int
    totalPages: int


def get_report_service() -> ReportService:
    return ReportService(report_store, catalog_store, internal_data_store, monitoring_store, recommendation_store)


def to_response(report: Report, version: ReportVersion) -> ReportResponse:
    return ReportResponse(
        id=report.id, title=report.title, reportPeriod=report.report_period.value,
        inputMode=report.input_mode, periodStart=report.period_start, periodEnd=report.period_end,
        status=report.status.value, inputSnapshotDates=list(report.input_snapshot_dates),
        approvedBy=report.approved_by, approvedAt=report.approved_at,
        currentVersion=VersionResponse(
            id=version.id, version=version.version, markdown=version.markdown,
            contentDigest=version.content_digest, changeSource=version.change_source,
            createdBy=version.created_by, createdAt=version.created_at,
        ),
        createdAt=report.created_at, updatedAt=report.updated_at,
    )


@router.post("", response_model=ReportResponse, status_code=status.HTTP_201_CREATED)
def create_report(payload: CreateReportRequest, user: User = Depends(get_current_user), service: ReportService = Depends(get_report_service)) -> ReportResponse:
    try:
        report, version = service.create(user.workspace_id, user.id, payload.title, payload.reportPeriod, payload.periodStart, payload.periodEnd)
    except PeriodInputIncompleteError as exc:
        raise api_error(409, "PERIOD_INPUT_INCOMPLETE", str(exc), {"missingDates": [d.isoformat() for d in exc.missing_dates]}) from exc
    except DuplicateDailyReportError as exc:
        raise api_error(409, "DAILY_REPORT_EXISTS", str(exc)) from exc
    except ValueError as exc:
        raise api_error(422, "REPORT_PERIOD_INVALID", str(exc)) from exc
    return to_response(report, version)


@router.get("", response_model=ReportListResponse)
def list_reports(
    page: int = 1,
    pageSize: int = 20,
    user: User = Depends(get_current_user),
    service: ReportService = Depends(get_report_service),
) -> ReportListResponse:
    if page < 1 or pageSize < 1 or pageSize > 100:
        raise api_error(422, "PAGINATION_INVALID", "page must be positive and pageSize must be between 1 and 100.")
    all_reports = service.store.list(user.workspace_id)
    start = (page - 1) * pageSize
    records = []
    for report in all_reports[start : start + pageSize]:
        version = service.store.get_version(report.current_version_id)
        if version:
            records.append(to_response(report, version))
    total = len(all_reports)
    return ReportListResponse(
        data=records,
        pagination=Pagination(
            page=page,
            pageSize=pageSize,
            totalItems=total,
            totalPages=(total + pageSize - 1) // pageSize if total else 0,
        ),
    )


@router.get("/{report_id}", response_model=ReportResponse)
def get_report(report_id: str, user: User = Depends(get_current_user), service: ReportService = Depends(get_report_service)) -> ReportResponse:
    try:
        return to_response(*service.get(user.workspace_id, report_id))
    except LookupError as exc:
        raise api_error(404, "REPORT_NOT_FOUND", "Report was not found.") from exc


@router.post("/{report_id}/versions", response_model=ReportResponse, status_code=status.HTTP_201_CREATED)
def save_report_version(report_id: str, payload: SaveVersionRequest, user: User = Depends(get_current_user), service: ReportService = Depends(get_report_service)) -> ReportResponse:
    try:
        return to_response(*service.save_version(user.workspace_id, report_id, user.id, payload.markdown))
    except LookupError as exc:
        raise api_error(404, "REPORT_NOT_FOUND", "Report was not found.") from exc
    except ValueError as exc:
        raise api_error(409, "REPORT_ALREADY_APPROVED", str(exc)) from exc


@router.post("/{report_id}/approve", response_model=ReportResponse)
def approve_report(report_id: str, user: User = Depends(get_current_user), service: ReportService = Depends(get_report_service)) -> ReportResponse:
    try:
        return to_response(*service.approve(user.workspace_id, report_id, user.id))
    except LookupError as exc:
        raise api_error(404, "REPORT_NOT_FOUND", "Report was not found.") from exc


def approved_export(report_id: str, user: User, service: ReportService) -> tuple[Report, ReportVersion]:
    try:
        report, version = service.get(user.workspace_id, report_id)
    except LookupError as exc:
        raise api_error(404, "REPORT_NOT_FOUND", "Report was not found.") from exc
    if report.status.value != "APPROVED":
        raise api_error(409, "REPORT_NOT_APPROVED", "Report must be approved before export.")
    return report, version


@router.get("/{report_id}/exports/markdown")
def export_markdown(report_id: str, user: User = Depends(get_current_user), service: ReportService = Depends(get_report_service)) -> Response:
    report, version = approved_export(report_id, user, service)
    return Response(content=version.markdown.encode(), media_type="text/markdown; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{report.id}.md"'})


@router.get("/{report_id}/exports/docx")
def export_docx(report_id: str, user: User = Depends(get_current_user), service: ReportService = Depends(get_report_service)) -> StreamingResponse:
    report, version = approved_export(report_id, user, service)
    document = Document()
    for line in version.markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            document.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line:
            document.add_paragraph(line)
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{report.id}.docx"'})
